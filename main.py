from crewai import Agent, Task, Crew, LLM
from crewai_tools import SerperDevTool
import streamlit as st
import re
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
load_dotenv()

# Streamlit page config
st.set_page_config(page_title="AI Blog Generator", page_icon="📰", layout="wide")

# Title and description
st.title("🤖 AI Blog Generator")
st.markdown("Generate comprehensive blog posts about any topic using AI agents.")

# Sidebar - Only Advanced Settings and How to use
with st.sidebar:
    st.header("Settings")
    
    # Advanced Settings
    st.markdown("### Advanced Settings")
    temperature = st.slider("Temperature", 0.0, 1.0, 0.7)
    
    # Add helpful information
    with st.expander("How to use"):
        st.markdown("""
        1. Enter your desired topic in the text area in the main section
        2. Adjust the temperature if needed (higher = more creative)
        3. Click 'Generate Content' to start
        4. Wait for the AI to generate your article
        5. Download the result as a markdown / docx file
        """)

def generate_content(topic):
    llm = LLM(
        model="command-r",
        temperature=0.7
    )

    search_tool = SerperDevTool(n_results=10)

    # First Agent: Senior Research Analyst
    senior_research_analyst = Agent(
        role="Senior Research Analyst",
        goal=f"Research, analyze, and synthesize comprehensive information on {topic} from reliable web sources",
        backstory="You're an expert research analyst with advanced web research skills. "
                "You excel at finding, analyzing, and synthesizing information from "
                "across the internet using search tools. You're skilled at "
                "distinguishing reliable sources from unreliable ones, "
                "fact-checking, cross-referencing information, and "
                "identifying key patterns and insights. You provide "
                "well-organized research briefs with proper citations "
                "and source verification. Your analysis includes both "
                "raw data and interpreted insights, making complex "
                "information accessible and actionable.",
        allow_delegation=False,
        verbose=True,
        tools=[search_tool],
        llm=llm
    )

    # Second Agent: Content Writer
    content_writer = Agent(
        role="Content Writer",
        goal="Transform research findings into engaging blog posts while maintaining accuracy",
        backstory="You're a skilled content writer specialized in creating "
                "engaging, accessible content from technical research. "
                "You work closely with the Senior Research Analyst and excel at maintaining the perfect "
                "balance between informative and entertaining writing, "
                "while ensuring all facts and citations from the research "
                "are properly incorporated. You have a talent for making "
                "complex topics approachable without oversimplifying them.",
        allow_delegation=False,
        verbose=True,
        llm=llm
    )

    # Research Task
    research_task = Task(
        description=("""
            1. Conduct comprehensive research on {topic} including:
                - Recent developments and news
                - Key industry trends and innovations
                - Expert opinions and analyses
                - Statistical data and market insights
            2. Evaluate source credibility and fact-check all information
            3. Organize findings into a structured research brief
            4. Include all relevant citations and sources
        """),
        expected_output="""A detailed research report containing:
            - Executive summary of key findings
            - Comprehensive analysis of current trends and developments
            - List of verified facts and statistics
            - All citations and links to original sources
            - Clear categorization of main themes and patterns
            Please format with clear sections and bullet points for easy reference.""",
        agent=senior_research_analyst
    )

    # Writing Task
    writing_task = Task(
        description=("""
            Using the research brief provided, create an engaging blog post that:
            1. Transforms technical information into accessible content
            2. Maintains all factual accuracy and citations from the research
            3. Includes:
                - Attention-grabbing introduction
                - Well-structured body sections with clear headings
                - Compelling conclusion
            4. Preserves all source citations in [Source: URL] format
            5. Includes a References section at the end
        """),
        expected_output="""A polished blog post in markdown format that:
            - Engages readers while maintaining accuracy
            - Contains properly structured sections
            - Includes Inline citations hyperlinked to the original source url
            - Presents information in an accessible yet informative way
            - Follows proper markdown formatting, use H1 for the title and H3 for the sub-sections""",
        agent=content_writer
    )

    # Create Crew
    crew = Crew(
        agents=[senior_research_analyst, content_writer],
        tasks=[research_task, writing_task],
        verbose=True
    )

    return crew.kickoff(inputs={"topic": topic})

def sanitize_filename(topic):
    """Sanitize topic string to create a safe filename"""
    # Remove or replace problematic characters
    sanitized = re.sub(r'[^\w\s-]', '', topic.strip())  # Keep only alphanumeric, spaces, and hyphens
    sanitized = re.sub(r'[-\s]+', '_', sanitized)  # Replace spaces and multiple hyphens with single underscore
    sanitized = sanitized.lower()  # Convert to lowercase
    
    # Ensure it's not empty and not too long
    if not sanitized:
        sanitized = "article"
    elif len(sanitized) > 50:  # Limit length
        sanitized = sanitized[:50]
    
    return sanitized

def markdown_to_docx(markdown_content):
    """Convert markdown content to a DOCX document"""
    doc = Document()
    
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:  # Empty line
            continue  # Skip empty lines to reduce extra spacing
            
        # Handle horizontal separators - add a proper horizontal line
        if line.startswith('---') or line == '---':
            # Add a horizontal line using a paragraph with bottom border
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_after = Inches(0.1)
            p_format.space_before = Inches(0.1)
            # Add a simple line of dashes as a visual separator
            run = p.add_run('_' * 50)
            run.font.color.rgb = None  # Use default color
            continue
            
        if line.startswith('# '):  # H1 - Title
            title = line[2:].strip()
            # Remove markdown formatting from title
            title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        elif line.startswith('#### '):  # H4 - Subheading
            subtitle = line[5:].strip()
            # Remove markdown formatting from subtitle
            subtitle = re.sub(r'\*\*(.*?)\*\*', r'\1', subtitle)
            doc.add_heading(subtitle, level=3)
            
        elif line.startswith('### '):  # H3 - Subheading
            subtitle = line[4:].strip()
            # Remove markdown formatting from subtitle
            subtitle = re.sub(r'\*\*(.*?)\*\*', r'\1', subtitle)
            doc.add_heading(subtitle, level=2)
            
        elif line.startswith('## '):  # H2 - Subheading
            subtitle = line[3:].strip()
            # Remove markdown formatting from subtitle
            subtitle = re.sub(r'\*\*(.*?)\*\*', r'\1', subtitle)
            doc.add_heading(subtitle, level=2)
            
        elif line.startswith('- ') or line.startswith('* '):  # Bullet points
            bullet_text = line[2:].strip()
            paragraph = doc.add_paragraph(style='List Bullet')
            add_formatted_text_with_links(paragraph, bullet_text)
            
        elif line.startswith('1. ') or re.match(r'^\d+\. ', line):  # Numbered lists
            numbered_text = re.sub(r'^\d+\. ', '', line).strip()
            paragraph = doc.add_paragraph(style='List Number')
            add_formatted_text_with_links(paragraph, numbered_text)
            
        else:  # Regular paragraph
            if line and not line.startswith('#'):  # Make sure it's not an unhandled heading
                paragraph = doc.add_paragraph()
                add_formatted_text_with_links(paragraph, line)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def markdown_to_docx(markdown_content):
    """Convert markdown content to a DOCX document"""
    doc = Document()
    
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:  # Empty line
            continue  # Skip empty lines to reduce extra spacing
            
        # Handle horizontal separators - add a simple separator
        if line.startswith('---') or line == '---':
            p = doc.add_paragraph()
            p.add_run('─' * 50)  # Use a Unicode line character
            continue
            
        if line.startswith('# '):  # H1 - Title
            title = line[2:].strip()
            # Remove markdown formatting from title
            title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        elif line.startswith('#### '):  # H4 - Subheading
            subtitle = line[5:].strip()
            # Remove markdown formatting from subtitle
            subtitle = re.sub(r'\*\*(.*?)\*\*', r'\1', subtitle)
            doc.add_heading(subtitle, level=3)
            
        elif line.startswith('### '):  # H3 - Subheading
            subtitle = line[4:].strip()
            # Remove markdown formatting from subtitle
            subtitle = re.sub(r'\*\*(.*?)\*\*', r'\1', subtitle)
            doc.add_heading(subtitle, level=2)
            
        elif line.startswith('## '):  # H2 - Subheading
            subtitle = line[3:].strip()
            # Remove markdown formatting from subtitle
            subtitle = re.sub(r'\*\*(.*?)\*\*', r'\1', subtitle)
            doc.add_heading(subtitle, level=2)
            
        elif line.startswith('- ') or line.startswith('* '):  # Bullet points
            bullet_text = line[2:].strip()
            paragraph = doc.add_paragraph(style='List Bullet')
            add_simple_formatted_text(paragraph, bullet_text)
            
        elif line.startswith('1. ') or re.match(r'^\d+\. ', line):  # Numbered lists
            numbered_text = re.sub(r'^\d+\. ', '', line).strip()
            paragraph = doc.add_paragraph(style='List Number')
            add_simple_formatted_text(paragraph, numbered_text)
            
        else:  # Regular paragraph
            if line and not line.startswith('#'):  # Make sure it's not an unhandled heading
                paragraph = doc.add_paragraph()
                add_simple_formatted_text(paragraph, line)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def add_simple_formatted_text(paragraph, text):
    """Add formatted text with proper clickable hyperlinks"""
    
    # Pattern 1: Handle "Text" Link (URL) format - common in AI generated content
    link_with_url_pattern = r'Link\s+\(([^)]+)\)'
    
    # Pattern 2: Standard markdown links [text](url)
    markdown_link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    # Pattern 3: Plain URLs (but not those preceded by "Source:") - FIXED to capture full URL
    plain_url_pattern = r'(?<!Source:\s)(https?://[^\s\)\],]+)'
    
    # First, handle the "Link (URL)" format by replacing it with just the URL
    text = re.sub(link_with_url_pattern, r'\1', text)
    
    current_pos = 0
    processed_any_links = False
    
    # Handle markdown links [text](url)
    for match in re.finditer(markdown_link_pattern, text):
        processed_any_links = True
        
        # Add text before the link
        if match.start() > current_pos:
            before_text = text[current_pos:match.start()]
            add_text_formatting(paragraph, before_text)
        
        # Create proper hyperlink
        link_text = match.group(1)
        url = match.group(2)
        add_hyperlink(paragraph, url, link_text)
        
        current_pos = match.end()
    
    # If no markdown links were found, look for plain URLs (excluding Source: URLs)
    if not processed_any_links:
        url_current_pos = 0
        
        for url_match in re.finditer(plain_url_pattern, text):
            # Double-check if this URL is preceded by "Source:" - if so, skip making it a hyperlink
            start_check = max(0, url_match.start() - 15)  # Check 15 characters before to be safe
            preceding_text = text[start_check:url_match.start()]
            
            if "Source:" in preceding_text:
                # Don't create hyperlink for Source: URLs, just treat as regular text
                continue
            
            # Add text before the URL
            if url_match.start() > url_current_pos:
                before_url_text = text[url_current_pos:url_match.start()]
                add_text_formatting(paragraph, before_url_text)
            
            # Create hyperlink with full URL as both link and display text - FIXED
            url = url_match.group(0)  # Changed from group(1) to group(0) to get full match
            add_hyperlink(paragraph, url, url)
            
            url_current_pos = url_match.end()
        
        # Add any remaining text (including Source: URLs as regular text)
        if url_current_pos < len(text):
            remaining_text = text[url_current_pos:]
            add_text_formatting(paragraph, remaining_text)
        elif url_current_pos == 0:
            # No links found at all, just add formatted text
            add_text_formatting(paragraph, text)
    else:
        # Add any remaining text after markdown links
        if current_pos < len(text):
            remaining_text = text[current_pos:]
            add_text_formatting(paragraph, remaining_text)

def add_hyperlink(paragraph, url, text):
    """Add a proper hyperlink to a paragraph using python-docx"""
    try:
        # Get the document from paragraph
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # Create hyperlink element
        hyperlink = paragraph._element.makeelement(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink",
            {"{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id": r_id}
        )
        
        # Create run element for the hyperlink text
        run = hyperlink.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        
        # Create run properties for hyperlink styling
        rPr = run.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
        
        # Set color to blue
        color = rPr.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color", 
                               {"{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "0563C1"})
        rPr.append(color)
        
        # Set underline
        underline = rPr.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u", 
                                   {"{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "single"})
        rPr.append(underline)
        
        run.append(rPr)
        
        # Create text element
        text_elem = run.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        text_elem.text = text
        run.append(text_elem)
        
        hyperlink.append(run)
        paragraph._element.append(hyperlink)
        
    except Exception as e:
        # Fallback: if hyperlink creation fails, add as regular underlined text
        print(f"Hyperlink creation failed: {e}")
        run = paragraph.add_run(text)
        run.underline = True

def add_text_formatting(paragraph, text):
    """Add text with bold and italic formatting - simplified version"""
    if not text.strip():
        return
    
    # Simple approach: handle bold and italic sequentially to avoid conflicts
    current_text = text
    
    # Handle bold text **text**
    bold_pattern = r'\*\*(.*?)\*\*'
    parts = re.split(bold_pattern, current_text)
    
    for i, part in enumerate(parts):
        if i % 2 == 0:  # Regular text
            # Handle italic in regular text
            italic_pattern = r'(?<!\*)\*([^*]+)\*(?!\*)'
            italic_parts = re.split(italic_pattern, part)
            
            for j, italic_part in enumerate(italic_parts):
                if j % 2 == 0:  # Regular text
                    if italic_part:
                        paragraph.add_run(italic_part)
                else:  # Italic text
                    if italic_part:
                        run = paragraph.add_run(italic_part)
                        run.italic = True
        else:  # Bold text
            if part:
                run = paragraph.add_run(part)
                run.bold = True

# Main content area - Topic input and Generate button moved here
st.markdown("### Enter Your Topic")

# Topic input in main area
topic = st.text_area(
    "What would you like to write about?",
    height=100,
    placeholder="Enter the topic you want to generate content about...",
    label_visibility="collapsed"
)

# Generate button in main area
generate_button = st.button("Generate Content", type="primary", use_container_width=False)

# Content generation and display
if generate_button:
    if topic.strip():
        with st.spinner('Generating content... This may take a moment.'):
            try:
                result = generate_content(topic)
                st.markdown(result)
                
                # Add download buttons with sanitized filenames
                base_filename = sanitize_filename(topic)
                
                # Create two columns for download buttons
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        label="📄 Download as Markdown",
                        data=result.raw,
                        file_name=f"{base_filename}_article.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
                
                with col2:
                    try:
                        # Generate DOCX content
                        docx_content = markdown_to_docx(str(result.raw))
                        st.download_button(
                            label="📝 Download as Word Document",
                            data=docx_content,
                            file_name=f"{base_filename}_article.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as docx_error:
                        st.error(f"Error generating DOCX: {str(docx_error)}")
                        st.info("Markdown download is still available above.")
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
    else:
        st.warning("Please enter a topic before generating content.")